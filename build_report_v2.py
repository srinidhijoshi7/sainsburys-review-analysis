"""
build_report_v2.py — Final report assembler with all 9 requested changes:
  1. Cover sheet as page 1 (copied from Individual_Cover_Sheet.docx)
  2. Title page as page 2
  3. Abstract as page 3
  4. Table of Contents as page 4 (with dot leaders, styled like sample)
  5. Introduction starts page 5, page numbering restarts at 1 there
  6. Summary results table (BERT validation metrics + event impact)
  7. GitHub link in Appendix opening
  8. Interactive LDA dashboard mentioned + screenshot placeholder in Appendix C
  9. References as bullet list

Run from project root:
    python3 build_report_v2.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from pathlib import Path
import copy

PROJECT_ROOT = Path(__file__).parent
FIGS_DIR = PROJECT_ROOT / "figures"
COVER_SHEET_PATH = PROJECT_ROOT / "Individual_Cover_Sheet.docx"
OUTPUT_FILE = PROJECT_ROOT / "Sainsburys_Review_Analysis_Report_FINAL.docx"

GITHUB_PLACEHOLDER = "https://github.com/srinidhijoshi7/sainsburys-review-analysis"
INTERACTIVE_VIZ_PATH = "figures/lda_interactive_enhanced.html"

SCRIPTS_TO_APPEND = [
    "discover_products.py", "scrape_reviews.py", "clean_reviews.py",
    "sentiment_analysis.py", "apply_bert_full.py", "topic_modelling.py",
    "temporal_analysis.py", "build_interactive_page.py",
]

MAIN_FIGURES = {
    1:  ("fig01_rating_distribution.png",
         "Distribution of star ratings across the cleaned corpus (n = 64,916). "
         "The J-shaped distribution is characteristic of voluntary online review "
         "platforms (Hu et al., 2009)."),
    6:  ("fig06_confusion_matrix.png",
         "Confusion matrix for the RoBERTa sentiment classifier benchmarked "
         "on a stratified sample of 5,000 reviews (1,000 per star rating). "
         "Ground truth: 1–2 stars = negative, 3 = neutral, 4–5 = positive. "
         "Accuracy = 76.8%, Macro-F1 = 0.630."),
    11: ("fig11_topic_by_category.png",
         "Distribution of LDA complaint topics across product categories "
         "(row-normalised). Darker cells indicate stronger concentration."),
    12: ("fig12_top_words_per_topic.png",
         "Top 10 words per LDA topic (K = 5, c_v = 0.40), fitted on 5,257 "
         "high-confidence negative reviews. Labels assigned after inspection "
         "of top words and exemplar reviews."),
    13: ("fig13_sentiment_with_events.png",
         "Monthly mean BERT sentiment (top) and share negative (bottom), "
         "2023–2026, annotated with verified UK retail and macroeconomic events."),
    15: ("fig15_category_sentiment_heatmap.png",
         "Mean BERT sentiment by category and quarter. Categories sorted "
         "worst (top) to best (bottom). Red = negative; green = positive."),
}

APPENDIX_FIGURES = {
    "A1": ("fig02_reviews_over_time.png",
           "Review volume over time. The step increase in mid-2024 coincides "
           "with Sainsbury's INSTORE_REVIEWS campaign expansion."),
    "A2": ("fig03_category_counts.png",
           "Number of cleaned reviews per product category."),
    "A3": ("fig04_text_length.png",
           "Distribution of review length in words (clipped at 99th percentile)."),
    "A4": ("fig05_avg_rating_by_category.png",
           "Mean star rating by category. Categories below the overall mean "
           "(4.44) are shown in orange."),
    "A5": ("fig07_bert_vs_rating_heatmap.png",
           "Distribution of BERT predictions by star rating. The 3-star row "
           "splits between negative and positive — evidence that 'neutral' "
           "is not a natural language sentiment category."),
    "A6": ("fig08_sentiment_over_time.png",
           "Mean monthly BERT sentiment vs. star-rating ground truth on the "
           "5,000-review stratified sample."),
    "A7": ("fig09_lda_coherence.png",
           "LDA coherence (c_v) across K ∈ {5, 7, 10, 12}. K = 5 selected."),
    "A8": ("fig10_topic_prevalence_over_time.png",
           "Proportion of negative reviews by topic, by quarter."),
    "A9": ("fig14_topic_volume_with_events.png",
           "Absolute volume of negative reviews by topic, by month."),
}

REFERENCES = [
    "Bird, S., Klein, E. and Loper, E. (2009) Natural Language Processing "
    "with Python. Sebastopol, CA: O'Reilly Media.",
    "Blei, D.M., Ng, A.Y. and Jordan, M.I. (2003) 'Latent Dirichlet "
    "allocation', Journal of Machine Learning Research, 3, pp. 993–1022.",
    "British Psychological Society (2021) Ethics guidelines for internet-"
    "mediated research. Leicester: BPS.",
    "Choudhury, P., Wang, D., Carlson, N.A. and Khanna, T. (2019) 'Machine "
    "learning approaches to facial and text analysis: Discovering CEO oral "
    "communication styles', Strategic Management Journal, 40(11), pp. 1705–1732.",
    "Devlin, J., Chang, M.-W., Lee, K. and Toutanova, K. (2019) 'BERT: "
    "Pre-training of deep bidirectional transformers for language "
    "understanding', in Proceedings of NAACL-HLT 2019, pp. 4171–4186.",
    "Dorotic, M., Bijmolt, T.H.A. and Verhoef, P.C. (2012) 'Loyalty "
    "programmes: Current knowledge and research directions', International "
    "Journal of Management Reviews, 14(3), pp. 217–237.",
    "Hu, N., Zhang, J. and Pavlou, P.A. (2009) 'Overcoming the J-shaped "
    "distribution of product reviews', Communications of the ACM, 52(10), "
    "pp. 144–147.",
    "Humphreys, A. and Wang, R.J.-H. (2018) 'Automated text analysis for "
    "consumer research', Journal of Consumer Research, 44(6), pp. 1274–1306.",
    "Hunter, J.D. (2007) 'Matplotlib: A 2D graphics environment', "
    "Computing in Science & Engineering, 9(3), pp. 90–95.",
    "Hutto, C.J. and Gilbert, E. (2014) 'VADER: A parsimonious rule-based "
    "model for sentiment analysis of social media text', in Proceedings of "
    "ICWSM 2014, Ann Arbor, MI.",
    "Lamey, L., Deleersnyder, B., Dekimpe, M.G. and Steenkamp, J.-B.E.M. "
    "(2007) 'How business cycles contribute to private-label success', "
    "Journal of Marketing, 71(1), pp. 1–15.",
    "Liu, B. (2019) Sentiment Analysis: Mining opinions, sentiments, and "
    "emotions. 2nd edn. Cambridge: Cambridge University Press.",
    "Loureiro, D. et al. (2022) 'TimeLMs: Diachronic language models from "
    "Twitter', in Proceedings of ACL 2022: System Demonstrations, pp. 251–260.",
    "Mabey, B. (2018) pyLDAvis (Version 3.4.0) [Computer program]. Available "
    "at: https://github.com/bmabey/pyLDAvis (Downloaded: 23 April 2026).",
    "McKinney, W. (2010) 'Data structures for statistical computing in Python', "
    "in Proceedings of the 9th Python in Science Conference, pp. 51–56.",
    "Office for National Statistics (2023) Consumer price inflation, UK: "
    "March 2023. Newport: ONS. Available at: https://www.ons.gov.uk/economy/"
    "inflationandpriceindices/bulletins/consumerpriceinflation/march2023 "
    "(Accessed: 23 April 2026).",
    "Pang, B. and Lee, L. (2008) 'Opinion mining and sentiment analysis', "
    "Foundations and Trends in Information Retrieval, 2(1–2), pp. 1–135.",
    "Paulhus, D.L. (1991) 'Measurement and control of response bias', in "
    "Robinson, J.P. et al. (eds.) Measures of Personality and Social "
    "Psychological Attitudes. San Diego: Academic Press, pp. 17–59.",
    "Python Software Foundation (2024) Python (Version 3.9.6) [Computer "
    "program]. Available at: https://www.python.org (Downloaded: 23 April 2026).",
    "Richardson, L. (2024) Beautiful Soup (Version 4.12) [Computer program]. "
    "Available at: https://www.crummy.com/software/BeautifulSoup/ "
    "(Downloaded: 23 April 2026).",
    "Röder, M., Both, A. and Hinneburg, A. (2015) 'Exploring the space of "
    "topic coherence measures', in Proceedings of WSDM 2015, pp. 399–408.",
    "Řehůřek, R. and Sojka, P. (2010) 'Software framework for topic modelling "
    "with large corpora', in Proceedings of LREC 2010 NLP Frameworks Workshop, "
    "pp. 45–50.",
    "Schoenmueller, V., Netzer, O. and Stahl, F. (2020) 'The polarity of "
    "online reviews', Journal of Marketing Research, 57(5), pp. 853–877.",
    "Sharp, B. and Sharp, A. (1997) 'Loyalty programs and their impact on "
    "repeat-purchase loyalty patterns', International Journal of Research in "
    "Marketing, 14(5), pp. 473–486.",
    "Sia, S., Dalmia, A. and Mielke, S.J. (2020) 'Tired of topic models? ',"
    "in Proceedings of EMNLP 2020, pp. 1728–1736.",
    "Waskom, M.L. (2021) 'Seaborn: statistical data visualization', "
    "Journal of Open Source Software, 6(60), p. 3021.",
    "Wolf, T. et al. (2020) 'Transformers: State-of-the-art natural language "
    "processing', in Proceedings of EMNLP 2020 System Demonstrations, pp. 38–45.",
    "Wu, D. (2023) 'Text-based measure of supply chain risk exposure', "
    "Management Science, forthcoming.",
]


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_section_break_next_page(doc):
    """Adds a next-page section break, used to restart page numbering."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    sectPr = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "12240")
    pgSz.set(qn("w:h"), "15840")
    sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    for attr, val in [("w:top","1440"),("w:right","1440"),
                      ("w:bottom","1440"),("w:left","1440"),
                      ("w:header","720"),("w:footer","720"),("w:gutter","0")]:
        pgMar.set(qn(attr), val)
    sectPr.append(pgMar)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")  # restart numbering at 1
    sectPr.append(pgNumType)
    p._p.append(sectPr)


def set_footer_page_number(section, start=None):
    """Add centred page number to footer of a section."""
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if start is not None:
        # Set section start page
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn("w:pgNumType"))
        if pgNumType is None:
            pgNumType = OxmlElement("w:pgNumType")
            sectPr.append(pgNumType)
        pgNumType.set(qn("w:start"), str(start))

    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_heading(doc, text, level=1, number=None):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(f"{number + '  ' if number else ''}{text}")
    run.font.name = "Calibri"
    run.font.size = Pt(16 if level == 1 else (13 if level == 2 else 11))
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def add_rich(doc, segments):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    for seg in segments:
        run = p.add_run(seg["text"])
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.italic = seg.get("italic", False)
        run.bold = seg.get("bold", False)
    return p


def add_figure(doc, fig_number, fig_path, caption, width_inches=6.0):
    if not fig_path.exists():
        p = doc.add_paragraph(f"[MISSING: {fig_path.name}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(180, 0, 0)
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(fig_path), width=Inches(width_inches))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r1 = p_cap.add_run(f"Figure {fig_number}. ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Calibri"
    r2 = p_cap.add_run(caption)
    r2.italic = True; r2.font.size = Pt(10); r2.font.name = "Calibri"


def add_code_block(doc, code_text, filename=None):
    if filename:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"File: {filename}")
        r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(11)
    for line in code_text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        leading = len(line) - len(line.lstrip(" "))
        body = line.lstrip(" ")
        if leading:
            rs = p.add_run("\u00a0" * leading)
            rs.font.name = "Consolas"; rs.font.size = Pt(8)
        r = p.add_run(body or " ")
        r.font.name = "Consolas"; r.font.size = Pt(8)


# ---------------------------------------------------------------------------
# 1. Cover sheet — copy entire content from the provided docx
# ---------------------------------------------------------------------------

def build_cover_sheet(doc):
    """Copy cover sheet content into the document."""
    if not COVER_SHEET_PATH.exists():
        add_body(doc, "[Cover sheet not found — insert Individual_Cover_Sheet.docx here]")
        add_page_break(doc)
        return

    src = Document(COVER_SHEET_PATH)
    for element in src.element.body:
        # Deep copy each element from the source doc into ours
        imported = copy.deepcopy(element)
        doc.element.body.append(imported)

    add_page_break(doc)


# ---------------------------------------------------------------------------
# 2. Title page
# ---------------------------------------------------------------------------

def build_title_page(doc):
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Decoding the Voice of the Customer:")
    r.font.name = "Calibri"; r.font.size = Pt(24); r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "A Web-Scraping and NLP Analysis of Sainsbury's Customer Reviews\n"
        "through the UK Cost-of-Living Crisis, 2023–2026"
    )
    r.font.name = "Calibri"; r.font.size = Pt(14); r.italic = True

    for _ in range(5):
        doc.add_paragraph()

    for text, size, bold in [
        ("Srinidhi Joshi", 13, True),
        ("MSc Business Analytics", 12, False),
        ("University of Bristol", 12, False),
        ("", 11, False),
        ("Unit: Social Media and Web Analytics", 11, False),
        ("Submission: 30 April 2026", 11, False),
        ("Word count (main body): ~3,000", 11, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(size); r.bold = bold

    add_page_break(doc)


# ---------------------------------------------------------------------------
# 3. Abstract
# ---------------------------------------------------------------------------

def build_abstract(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Abstract")
    r.font.name = "Calibri"; r.font.size = Pt(16); r.bold = True

    doc.add_paragraph()  # spacer

    abstract_text = (
        "This report applies a two-stage natural language processing pipeline to "
        "64,916 Sainsbury's customer product reviews collected via web scraping "
        "from sainsburys.co.uk, spanning January 2023 to April 2026 — a period "
        "encompassing the peak of UK food inflation (19.1% in March 2023) and "
        "Sainsbury's subsequent strategic responses, including the launch of the "
        "Nectar Prices loyalty programme. A RoBERTa transformer model, validated "
        "against star-rating ground truth (accuracy = 76.8%, macro-F1 = 0.630 on "
        "a stratified 5,000-review sample), classifies sentiment at scale. Latent "
        "Dirichlet Allocation is then applied to 5,257 high-confidence negative "
        "reviews, revealing five coherent complaint themes: meat and protein quality "
        "failures, product reformulations, taste and texture disappointment, "
        "price-value and shrinkflation concerns, and prepared food underperformance. "
        "An event-annotated temporal analysis demonstrates that the largest positive "
        "sentiment shift in the dataset (Δ = +0.257 in monthly mean sentiment) "
        "coincided with the Nectar Prices launch, while the March 2024 IT outage "
        "produced no measurable product-review sentiment change — a finding that "
        "illuminates the boundary conditions of review-based analytics. Three "
        "actionable recommendations are developed: prioritising protein-category "
        "quality assurance, repositioning the Taste the Difference own-brand with "
        "value reassurance, and routing operational incident monitoring away from "
        "product-review data. The full pipeline — discovery scraper, review "
        "harvester, BERT classifier, LDA topic model, and temporal analysis — is "
        "openly available on GitHub."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    r = p.add_run(abstract_text)
    r.font.name = "Calibri"; r.font.size = Pt(11)

    doc.add_paragraph()

    # Keywords
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    r1 = p.add_run("Keywords: ")
    r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(11)
    r2 = p.add_run(
        "web scraping, sentiment analysis, BERT, RoBERTa, LDA topic modelling, "
        "customer reviews, UK grocery retail, Sainsbury's, cost-of-living crisis, "
        "loyalty programmes, natural language processing"
    )
    r2.italic = True; r2.font.name = "Calibri"; r2.font.size = Pt(11)

    add_page_break(doc)


# ---------------------------------------------------------------------------
# 4. Table of Contents (styled like the sample image — dot leaders)
# ---------------------------------------------------------------------------

def add_toc_entry(doc, number, title, page, level=1):
    """Add one TOC line with a tab-leader dot."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4 if level == 1 else 2)
    if level == 2:
        p.paragraph_format.left_indent = Cm(1.0)

    # Tab stop at right margin (page width 12240 - 2880 margins = 9360 twips)
    tab_stop = p.paragraph_format.tab_stops
    tab_stop.add_tab_stop(Twips(8640), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    label = f"{number}. {title}" if number else title
    r1 = p.add_run(label + "\t" + str(page))
    r1.font.name = "Calibri"
    r1.font.size = Pt(11 if level == 1 else 10)
    r1.bold = (level == 1)
    return p


def build_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Table of Contents")
    r.font.name = "Calibri"; r.font.size = Pt(16); r.bold = True

    doc.add_paragraph()

    # TOC entries — page numbers as they'll approximately appear
    # (User should update these manually or via Word's Update Field if needed)
    entries = [
        (1,  "Introduction",                               "1",  1),
        (2,  "Literature Review",                          "2",  1),
        (3,  "Methodology and Data",                       "3",  1),
        (None, "3.1  Research design and dataset",         "3",  2),
        (None, "3.2  Data collection pipeline",            "4",  2),
        (None, "3.3  Cleaning and preparation",            "5",  2),
        (None, "3.4  Analytical techniques",               "5",  2),
        (4,  "Analysis and Results",                       "7",  1),
        (None, "4.1  Rating distribution and validation",  "7",  1),
        (None, "4.2  Complaint themes",                    "8",  2),
        (None, "4.3  Temporal sentiment and events",       "9",  2),
        (5,  "Conclusion, Recommendations and Limitations","11", 1),
        (None, "5.1  Summary of findings",                 "11", 2),
        (None, "5.2  Recommendations",                     "11", 2),
        (None, "5.3  Limitations",                         "12", 2),
        (None, "References",                               "13", 1),
        (None, "Appendix A — Additional Figures",          "15", 1),
        (None, "Appendix B — Interactive LDA Dashboard",   "16", 1),
        (None, "Appendix C — Python Source Code",          "17", 1),
    ]

    for number, title, page, level in entries:
        add_toc_entry(doc, number, title, page, level)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(
        "Note: Page numbers are approximate. In Word, right-click and select "
        "'Update Field' on any TOC entry to refresh automatically."
    )
    r.font.name = "Calibri"; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(120, 120, 120)

    add_page_break(doc)


# ---------------------------------------------------------------------------
# 6. Tables (BERT validation + event impact)
# ---------------------------------------------------------------------------

def add_table(doc, title, headers, rows, table_number):
    """Add a numbered, captioned table."""
    # Caption above
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r1 = p.add_run(f"Table {table_number}. ")
    r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(10)
    r2 = p.add_run(title)
    r2.italic = True; r2.font.name = "Calibri"; r2.font.size = Pt(10)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    col_count = len(headers)
    content_width = 8640  # twips
    col_width = content_width // col_count

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Calibri"
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        # Grey shading on header
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E0E0E0")
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            row_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[c_idx].paragraphs[0].runs[0].font.name = "Calibri"
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Body sections
# ---------------------------------------------------------------------------

def build_introduction(doc):
    add_heading(doc, "Introduction", level=1, number="1.")
    add_body(doc,
        "The rise of e-commerce and social media has transformed how retailers "
        "understand their customers. Where traditional consumer research relies "
        "on surveys, focus groups and panel data — methods that are expensive, "
        "slow, and prone to social desirability bias (Paulhus, 1991) — modern "
        "analytics can draw on millions of unsolicited, time-stamped customer "
        "texts available at near-zero marginal cost (Humphreys and Wang, 2018). "
        "For UK grocery retailers operating in a post-pandemic, inflationary "
        "environment, the ability to detect and interpret customer sentiment in "
        "real time is no longer a competitive advantage but a requirement."
    )
    add_body(doc,
        "This report applies web scraping, transformer-based sentiment "
        "classification and topic modelling to 64,916 customer reviews from "
        "sainsburys.co.uk, published between January 2023 and April 2026. The "
        "period spans the peak of UK food inflation at 19.1% in March 2023, the "
        "highest rate in 45 years (Office for National Statistics, 2023), and "
        "Sainsbury's major strategic responses, most notably the launch of "
        "Nectar Prices on 11 April 2023. The research question is: how has "
        "customer sentiment toward Sainsbury's products evolved through the "
        "cost-of-living crisis, what are customers actually complaining about, "
        "and how should Sainsbury's customer experience strategy respond?"
    )
    add_body(doc,
        "The contribution is both methodological and practical. Methodologically, "
        "the study demonstrates a two-stage NLP pipeline — a RoBERTa sentiment "
        "filter followed by LDA topic modelling on a filtered negative sub-corpus "
        "— and validates the transformer's predictions against star-rating ground "
        "truth. Practically, the analysis identifies five distinct complaint "
        "themes, maps them by product category, and links a monthly sentiment "
        "timeline to verified UK retail events, yielding actionable, category-"
        "specific recommendations for Sainsbury's customer experience team "
        "(Section 5). The full analysis pipeline is publicly available on GitHub "
        f"({GITHUB_PLACEHOLDER})."
    )


def build_literature_review(doc):
    add_heading(doc, "Literature Review", level=1, number="2.")
    add_body(doc,
        "This study draws on three converging bodies of literature: (i) the "
        "rise of digital consumer-generated content as a substitute for "
        "traditional consumer research, (ii) methodological advances in NLP "
        "for business research, and (iii) retail-management work on loyalty "
        "programmes and price-sensitive customer behaviour."
    )
    add_rich(doc, [
        {"text": "Consumer-generated content and self-selection bias. ", "bold": True},
        {"text":
            "Online reviews offer scale and timeliness that survey methods "
            "cannot match (Humphreys and Wang, 2018). However, they are not "
            "representative. Hu "},
        {"text": "et al.", "italic": True},
        {"text": " (2009), in "},
        {"text": "Communications of the ACM", "italic": True},
        {"text":
            ", documented the 'J-shaped' distribution of online reviews, "
            "driven by self-selection. Schoenmueller, Netzer and Stahl (2020) "
            "in the "},
        {"text": "Journal of Marketing Research", "italic": True},
        {"text":
            " show this bias varies by category and platform. For this study, "
            "the skewed distribution (72.2% five-star) confirms the pattern "
            "and motivates concentrating topic modelling on negative reviews."},
    ])
    add_rich(doc, [
        {"text": "NLP methods in business research. ", "bold": True},
        {"text":
            "Lexicon-based approaches such as VADER (Hutto and Gilbert, 2014) "
            "struggle with negation and informal language. Transformer "
            "architectures introduced by Devlin "},
        {"text": "et al.", "italic": True},
        {"text":
            " (2019) with BERT represented a step change. Choudhury "},
        {"text": "et al.", "italic": True},
        {"text": " (2019) in "},
        {"text": "Strategic Management Journal", "italic": True},
        {"text":
            " demonstrate that ML text analysis can discover managerially "
            "meaningful patterns; Wu (2023) in "},
        {"text": "Management Science", "italic": True},
        {"text":
            " uses transformer text measures to quantify supply-chain risk "
            "exposure. For topic discovery, Blei, Ng and Jordan's (2003) LDA "
            "remains the dominant technique; Sia, Dalmia and Mielke (2020) "
            "show that filtering by sentiment prior to LDA produces more "
            "interpretable topics — the design adopted here."},
    ])
    add_rich(doc, [
        {"text": "Loyalty programmes and retail price sensitivity. ", "bold": True},
        {"text":
            "Sharp and Sharp (1997) in the "},
        {"text": "International Journal of Research in Marketing", "italic": True},
        {"text":
            " established scepticism about loyalty scheme retention effects; "
            "Dorotic, Bijmolt and Verhoef (2012) in "},
        {"text": "International Journal of Management Reviews", "italic": True},
        {"text":
            " show scheme effects on perceived value can be substantial, "
            "particularly during downturns. Lamey "},
        {"text": "et al.", "italic": True},
        {"text": " (2007) in the "},
        {"text": "Journal of Marketing", "italic": True},
        {"text":
            " document that premium private-label products are more vulnerable "
            "to recessionary demand shifts — a prediction confirmed by the "
            "Taste the Difference findings in Section 4.3."},
    ])


def build_methodology(doc):
    add_heading(doc, "Methodology and Data", level=1, number="3.")

    add_heading(doc, "Research design and dataset", level=2, number="3.1")
    add_body(doc,
        "This study uses web scraping and NLP methods to analyse customer "
        "reviews of Sainsbury's, the UK's second-largest supermarket. The "
        "reasoning for using digital review data rather than surveys is that "
        "surveys are expensive to run, often suffer from social desirability "
        "bias (Paulhus, 1991), and cannot easily capture customer experience "
        "at the scale needed to see patterns across thousands of products. "
        "Online reviews, by contrast, are written voluntarily, cover a wide "
        "product range, and come with star ratings that let us check our "
        "text-based sentiment estimates against a numerical ground truth "
        "(Hu et al., 2009)."
    )
    add_body(doc,
        "The analysis covers reviews published between 1 January 2023 and "
        "22 April 2026. This window was chosen because it includes the peak "
        "of UK food inflation in March 2023 at 19.1% — the highest annual "
        "rate since 1977 (Office for National Statistics, 2023) — and the "
        "period when Sainsbury's launched Nectar Prices (11 April 2023) and "
        "expanded its Aldi Price Match scheme."
    )

    add_heading(doc, "Data collection pipeline", level=2, number="3.2")
    add_body(doc,
        "Reviews were collected from sainsburys.co.uk, where verified "
        "customers leave feedback via the Bazaarvoice platform. Product-level "
        "reviews were chosen over Trustpilot or Twitter data because each "
        "review is tied to a specific product (enabling category segmentation), "
        "every review carries a star rating (enabling model validation), and "
        "most reviewers are actual buyers (reducing non-customer noise)."
    )
    add_body(doc,
        "A two-stage scraper was built in Python. Stage one queried "
        "Sainsbury's public search API across 25 category keywords to return "
        "2,500 unique products with reviews. Stage two pulled reviews from "
        "the Bazaarvoice endpoint in batches, capped at 50 reviews per "
        "product to prevent high-volume items (e.g. milk, 1,838 reviews) "
        "from dominating. The raw collection totalled 88,876 reviews. "
        "All scraping followed a polite rate limit of one request per 0.8 "
        "seconds; no authentication was bypassed; no personally identifiable "
        "information beyond self-chosen nicknames was retained. This follows "
        "ethical guidance for academic web scraping of public data "
        "(British Psychological Society, 2021)."
    )

    add_heading(doc, "Cleaning and preparation", level=2, number="3.3")
    add_body(doc,
        "The raw data was cleaned in four steps: (i) dropping three reviews "
        "with no text, (ii) filtering to the analysis window (removing 14,393 "
        "reviews), (iii) de-duplicating on review identifier (removing 9,564 "
        "Bazaarvoice-syndicated duplicates), and (iv) normalising whitespace "
        "and building a combined title+body text field. The final corpus "
        "contained 64,916 reviews from 2,644 unique products across 25 "
        "categories. Star ratings were collapsed to negative (1–2), neutral "
        "(3), and positive (4–5) for use as ground truth labels "
        "(Liu, 2019). Figure 1 shows the resulting rating distribution."
    )
    add_figure(doc, 1, FIGS_DIR / MAIN_FIGURES[1][0], MAIN_FIGURES[1][1])

    add_heading(doc, "Analytical techniques", level=2, number="3.4")
    add_rich(doc, [
        {"text": "Transformer sentiment classification. ", "bold": True},
        {"text":
            "The cardiffnlp/twitter-roberta-base-sentiment-latest model "
            "(Loureiro et al., 2022) was used in preference to VADER "
            "(Hutto and Gilbert, 2014) and generic BERT (Devlin et al., 2019) "
            "because its training data — approximately 124 million tweets — "
            "is closer in style to informal product reviews. A stratified "
            "sample of 5,000 reviews (1,000 per star rating) was classified "
            "and benchmarked against ground truth (Figure 6). The model was "
            "then applied to all 64,916 reviews using Apple MPS acceleration, "
            "completing in 15 minutes."},
    ])
    add_rich(doc, [
        {"text": "LDA topic modelling. ", "bold": True},
        {"text":
            "Latent Dirichlet Allocation (Blei, Ng and Jordan, 2003) was "
            "applied to 5,257 high-confidence negative reviews (BERT confidence "
            "≥ 0.80). After standard pre-processing (lowercasing, stopword "
            "removal, WordNet lemmatisation), the optimal K was selected via "
            "Röder et al.'s (2015) c_v coherence score across K ∈ {5,7,10,12}. "
            "K = 5 produced the best coherence (c_v = 0.40). An interactive "
            "pyLDAvis visualisation was generated and is available in "
            "Appendix B and on GitHub."},
    ])
    add_rich(doc, [
        {"text": "Event-annotated temporal analysis. ", "bold": True},
        {"text":
            "Monthly mean BERT sentiment and share-negative were plotted "
            "against seven verified UK events from ONS inflation bulletins "
            "and Sainsbury's corporate press releases. For each event, "
            "sentiment in the 30-day windows before and after was compared "
            "(Table 2). This analysis is associational, not causal."},
    ])


def build_analysis(doc):
    add_heading(doc, "Analysis and Results", level=1, number="4.")
    add_body(doc,
        "The analysis produced three sets of findings: (a) validation of the "
        "BERT classifier; (b) five distinct complaint themes and their category "
        "distribution; and (c) temporal links between sentiment shifts and "
        "verified UK retail events."
    )

    add_heading(doc, "Rating distribution and model validation", level=2, number="4.1")
    add_body(doc,
        "Figure 1 shows the rating distribution: 72.2% of reviews are 5-star "
        "and only 8.9% are 1–2 star. This 'J-shaped' curve reflects "
        "self-selection bias well-documented in the literature (Hu et al., "
        "2009; Schoenmueller, Netzer and Stahl, 2020), and implies that "
        "business-actionable signal is concentrated in the minority of "
        "negative reviews."
    )
    add_body(doc,
        "Table 1 presents the RoBERTa classifier performance on the 5,000-"
        "review stratified sample. The model achieved 76.8% accuracy and "
        "macro-F1 = 0.630. Figure 6 shows the confusion matrix. Recall was "
        "91% for negatives and 95% for positives, but only 12% for neutrals. "
        "Appendix Figure A5 shows that 3-star reviews ('neutral') in fact "
        "split 51%/38% between BERT-classified negative and positive text, "
        "confirming that 'neutral' is not a natural language sentiment "
        "(Pang and Lee, 2008). BERT's predictions are therefore more "
        "actionable than raw star counts."
    )

    # Table 1: BERT validation metrics
    add_table(doc,
        "RoBERTa sentiment classifier performance on stratified 5,000-review sample",
        ["Class", "Precision", "Recall", "F1-Score", "Support"],
        [
            ["Negative", "0.762", "0.912", "0.830", "2,000"],
            ["Neutral",  "0.486", "0.118", "0.190", "1,000"],
            ["Positive", "0.803", "0.948", "0.869", "2,000"],
            ["Macro avg","0.683", "0.659", "0.630", "5,000"],
            ["Accuracy", "—",     "—",     "0.768",  "5,000"],
        ],
        table_number=1
    )

    add_figure(doc, 6, FIGS_DIR / MAIN_FIGURES[6][0], MAIN_FIGURES[6][1])

    add_heading(doc, "Complaint themes", level=2, number="4.2")
    add_body(doc,
        "Applying LDA to the negative-review sub-corpus yielded five coherent "
        "themes (Figure 12). The c_v coherence peaked at K = 5 (c_v = 0.40) "
        "and the pyLDAvis distance map (Appendix B) shows all five topics "
        "occupying distinct quadrants, confirming they are genuinely separable."
    )

    # Bullet list of topics
    topics = [
        ("Meat/Protein Quality Failures (28%)",
         "tough, tasteless or inedible meat, chicken and fish; frequently "
         "described as waste of money. Dominates beef (73%), chicken (52%), "
         "salmon (40%)."),
        ("Reformulation / Recipe Changes (13%)",
         "anger at product formula changes ('the new recipe is awful'). "
         "Concentrated in shampoo (78%), washing powder (66%), coffee (49%)."),
        ("Taste & Texture Disappointment (18%)",
         "bland, dry, too sweet. Leads in bread (40%), sandwich (39%), "
         "cereal (35%)."),
        ("Price-Value & Shrinkflation (21%)",
         "price hikes, reduced pack weight, Nectar pricing frustration. "
         "Leads in toilet paper (69%), beer (66%), tea (59%)."),
        ("Prepared Food & Premium Disappointment (19%)",
         "pizza, ice cream, Taste the Difference underperforming expectations. "
         "Leads in pizza (49%), ice cream (48%), Taste the Difference (28%)."),
    ]
    for label, body in topics:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(label + " — ")
        r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(11)
        r2 = p.add_run(body)
        r2.font.name = "Calibri"; r2.font.size = Pt(11)

    add_figure(doc, 12, FIGS_DIR / MAIN_FIGURES[12][0], MAIN_FIGURES[12][1])
    add_figure(doc, 11, FIGS_DIR / MAIN_FIGURES[11][0], MAIN_FIGURES[11][1], 5.5)

    add_heading(doc, "Temporal sentiment and event linkage", level=2, number="4.3")
    add_body(doc,
        "Figure 13 plots monthly mean BERT sentiment and share negative across "
        "the full 40-month window. Table 2 summarises the 30-day pre/post "
        "sentiment comparison for each verified event."
    )

    # Table 2: Event impact
    add_table(doc,
        "30-day pre/post sentiment comparison around seven verified UK events "
        "(associational only; not causal)",
        ["Event", "Date", "Δ Sentiment", "% Neg before→after"],
        [
            ["UK food inflation peak (19.1%)", "Mar 2023", "+0.133", "22.7% → 16.3%"],
            ["Sainsbury's Nectar Prices launch", "Apr 2023", "+0.257", "23.9% → 11.7%"],
            ["Aldi Price Match expansion (400+)", "Aug 2023", "−0.023", "15.1% → 16.7%"],
            ["Sainsbury's IT outage", "Mar 2024", "−0.002", "15.7% → 16.0%"],
            ["Food inflation normalises (1.7%)", "May 2024", "+0.086", "14.0% → 9.2%"],
            ["Nectar Prices hits 5,000 products", "Sep 2024", "−0.035", "7.3% → 8.9%"],
            ["Aldi Price Match in Local stores", "Nov 2024", "+0.001", "9.8% → 10.1%"],
        ],
        table_number=2
    )

    add_body(doc,
        "Two findings stand out. First, the Nectar Prices launch coincided "
        "with the largest sentiment shift in the dataset (Δ = +0.257), with "
        "negative reviews halving from 23.9% to 11.7%. This is associational "
        "but the magnitude aligns with findings that price-based loyalty "
        "schemes shift customer affect in grocery retail (Dorotic, Bijmolt "
        "and Verhoef, 2012; Sharp and Sharp, 1997). Second, the IT outage "
        "produced effectively no change (Δ = −0.002), confirming that "
        "product-review sentiment reflects product experience rather than "
        "service incidents — a key limitation for practitioners."
    )

    add_figure(doc, 13, FIGS_DIR / MAIN_FIGURES[13][0], MAIN_FIGURES[13][1])

    add_body(doc,
        "Figure 15 provides a quarter-by-quarter sentiment heatmap per "
        "category. The worst-performing categories in early 2023 — beef, "
        "Taste the Difference, ready meal, juice, bread — all recovered by "
        "mid-2024. The broad recovery suggests macroeconomic factors drove "
        "improvement more than targeted product fixes. Notably, Taste the "
        "Difference showed sentiment of −0.50 in Q2 2023, consistent with "
        "Lamey et al.'s (2007) finding that premium private-label lines "
        "are recession-vulnerable."
    )

    add_figure(doc, 15, FIGS_DIR / MAIN_FIGURES[15][0], MAIN_FIGURES[15][1])


def build_conclusion(doc):
    add_heading(doc, "Conclusion, Recommendations and Limitations",
                level=1, number="5.")

    add_heading(doc, "Summary of findings", level=2, number="5.1")
    add_body(doc,
        "This study applied a two-stage NLP pipeline to 64,916 Sainsbury's "
        "customer reviews (January 2023 – April 2026). A validated RoBERTa "
        "sentiment classifier and LDA topic modelling on high-confidence "
        "negative reviews revealed five distinct complaint themes and a clear "
        "temporal narrative: sentiment tracked the UK food inflation cycle, "
        "with the largest positive shift coinciding with the Nectar Prices "
        "launch (Δ = +0.257), while the March 2024 IT outage produced no "
        "measurable product-review change."
    )

    add_heading(doc, "Recommendations", level=2, number="5.2")
    recs = [
        ("Prioritise protein-category quality assurance.",
         "Beef reviews are 73% Topic 0 (meat quality). Targeted supplier "
         "quality audits on fresh proteins offer the highest marginal return "
         "on customer-satisfaction investment (Figures 11 and 15)."),
        ("Reposition Taste the Difference with value reassurance.",
         "The premium own-brand had the deepest negative sentiment during "
         "the inflation peak (−0.50, Q2 2023). Strengthening value "
         "communication — pack-size transparency, ingredient provenance — "
         "would reduce this vulnerability in any future downturn "
         "(Lamey et al., 2007)."),
        ("Route operational incident monitoring away from product reviews.",
         "The IT outage produced no review-sentiment signal, confirming "
         "review data cannot detect service incidents. Twitter, customer-"
         "service transcripts and direct complaint channels should be used "
         "instead for operational monitoring."),
    ]
    for i, (title, body) in enumerate(recs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{i}. ").bold = True
        r_title = p.add_run(title + " ")
        r_title.bold = True; r_title.font.name = "Calibri"; r_title.font.size = Pt(11)
        r_body = p.add_run(body)
        r_body.font.name = "Calibri"; r_body.font.size = Pt(11)

    add_heading(doc, "Limitations", level=2, number="5.3")
    add_body(doc,
        "Three limitations should be noted. First, the 72% five-star "
        "concentration reflects self-selection bias (Hu et al., 2009) — "
        "averages understate dissatisfaction. Second, the pre-post event "
        "comparisons are associational, not causal. Third, the RoBERTa model "
        "was trained on Twitter data; domain adaptation to grocery text could "
        "improve performance, particularly on the neutral class (recall = 12%)."
    )


# ---------------------------------------------------------------------------
# References (bullet form as requested)
# ---------------------------------------------------------------------------

def build_references(doc):
    add_heading(doc, "References", level=1)
    add_body(doc,
        "All references are formatted in Harvard style. Software citations "
        "follow the BibGuru Harvard format for computer programmes."
    )
    doc.add_paragraph()
    for ref in sorted(REFERENCES):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(ref)
        run.font.name = "Calibri"
        run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Appendix A — Additional figures
# ---------------------------------------------------------------------------

def build_appendix_a(doc):
    add_heading(doc, "Appendix A — Additional Figures", level=1)
    add_body(doc,
        "Figures A1–A9 supplement the main-body figures. They include "
        "data-overview plots, additional model-validation charts, and "
        "topic-over-time visualisations referenced in the main text."
    )
    for label, (filename, caption) in APPENDIX_FIGURES.items():
        add_figure(doc, f"A{label[-1]}",
                   FIGS_DIR / filename, caption)


# ---------------------------------------------------------------------------
# Appendix B — Interactive LDA dashboard
# ---------------------------------------------------------------------------

def build_appendix_b_interactive(doc):
    add_heading(doc, "Appendix B — Interactive LDA Topic Dashboard", level=1)
    add_body(doc,
        "An interactive LDA visualisation was built using pyLDAvis "
        "(Mabey, 2018) and wrapped in a branded HTML page with explanatory "
        "panels for non-technical readers. The dashboard includes:"
    )
    # Bullet list of features
    features = [
        "An intertopic distance map (multidimensional scaling) showing the "
        "five complaint topics as circles, sized by prevalence and positioned "
        "by semantic distance — circles far apart are genuinely distinct.",
        "A per-topic word chart on the right, with a λ (lambda) relevance "
        "slider: λ near 0 shows the most distinctive words per topic; "
        "λ near 1 shows the most frequent words. Recommended setting: λ ≈ 0.6 "
        "(Sievert and Shirley, 2014).",
        "A topic legend mapping each numbered topic (1–5) to its human-"
        "readable label and a plain-English description of the dominant "
        "complaint pattern.",
        "A 'How to read this' guide for managers unfamiliar with LDA output.",
    ]
    for f in features:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f)
        r.font.name = "Calibri"; r.font.size = Pt(11)

    doc.add_paragraph()
    add_body(doc,
        f"The dashboard is included with the project source code and can be "
        f"opened locally by navigating to the project folder and opening "
        f"'{INTERACTIVE_VIZ_PATH}' in any web browser. It is also available "
        f"via the GitHub repository linked below."
    )

    # Screenshot placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("[Screenshot: lda_interactive_enhanced.html — "
                  "open the HTML file in Chrome/Safari to view]")
    r.italic = True; r.font.name = "Calibri"; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r1 = p_cap.add_run("Figure B1. ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Calibri"
    r2 = p_cap.add_run(
        "Interactive LDA dashboard (lda_interactive_enhanced.html). "
        "Built with pyLDAvis (Mabey, 2018); styled with custom CSS. "
        "Open in browser to explore topic-word distributions interactively."
    )
    r2.italic = True; r2.font.size = Pt(10); r2.font.name = "Calibri"

    # GitHub link block
    add_page_break(doc)
    add_heading(doc, "Appendix C — GitHub Repository", level=1)
    add_body(doc,
        "The complete analysis pipeline — all Python scripts, requirements, "
        "and the interactive HTML dashboard — is publicly available on GitHub:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(GITHUB_PLACEHOLDER)
    r.bold = True; r.font.name = "Consolas"; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 70, 180)

    add_body(doc,
        "The repository contains the following files:"
    )
    repo_files = [
        "discover_products.py — Stage 1: product discovery via Sainsbury's search API",
        "scrape_reviews.py — Stage 2: review harvesting via Bazaarvoice API",
        "clean_reviews.py — Stage 3: data cleaning, deduplication, feature engineering",
        "sentiment_analysis.py — Stage 4a: BERT validation on stratified sample",
        "apply_bert_full.py — Stage 4b: BERT applied to full 64,916-review corpus",
        "topic_modelling.py — Stage 5: LDA topic modelling with coherence selection",
        "temporal_analysis.py — Stage 6: event-annotated temporal analysis",
        "build_interactive_page.py — Stage 7: interactive pyLDAvis HTML dashboard",
        "figures/ — all 15 analysis figures",
        "data/ — products.csv, reviews_clean.csv, reviews_full_sentiment.csv, "
        "negative_reviews_with_topics.csv, monthly_sentiment.csv",
        "requirements.txt — full Python dependency list",
    ]
    for f in repo_files:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f)
        r.font.name = "Calibri"; r.font.size = Pt(10)

    add_body(doc,
        "To reproduce the analysis: clone the repository, install dependencies "
        "with pip install -r requirements.txt, and run the scripts in order "
        "(Stage 1 → 7). Runtime on an Apple Silicon MacBook is approximately "
        "90 minutes end-to-end (bulk of this is Stage 4b: 15 minutes, "
        "and Stage 2: 30 minutes)."
    )


# ---------------------------------------------------------------------------
# Appendix D — Python source code
# ---------------------------------------------------------------------------

def build_appendix_code(doc):
    add_heading(doc, "Appendix D — Python Source Code", level=1)
    add_body(doc,
        "The full Python source code is reproduced below. Each script is "
        "annotated with comments indicating its purpose, inputs and outputs. "
        "Scripts are listed in pipeline execution order. All software cited "
        "in the References per Harvard computer programme citation guidelines "
        "(BibGuru, 2025)."
    )
    for script_name in SCRIPTS_TO_APPEND:
        script_path = PROJECT_ROOT / script_name
        if not script_path.exists():
            add_body(doc, f"[Script not found: {script_name}]")
            continue
        add_code_block(doc, script_path.read_text(encoding="utf-8"),
                       filename=script_name)
        add_page_break(doc)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Page margins — all sections get 2.54cm (1 inch)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    print("Building cover sheet (page 1)...")
    build_cover_sheet(doc)

    print("Building title page (page 2)...")
    build_title_page(doc)

    print("Building abstract (page 3)...")
    build_abstract(doc)

    print("Building table of contents (page 4)...")
    build_toc(doc)

    # Page numbering restarts at 1 from here
    # We add a continuous section break then set the new section's footer
    print("Setting page number restart...")
    # Add a new section that starts on the next page with page number 1
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.top_margin = Cm(2.54)
    new_section.bottom_margin = Cm(2.54)
    new_section.left_margin = Cm(2.54)
    new_section.right_margin = Cm(2.54)
    set_footer_page_number(new_section, start=1)

    print("Building introduction (body page 1)...")
    build_introduction(doc)

    print("Building literature review...")
    build_literature_review(doc)

    print("Building methodology...")
    build_methodology(doc)

    print("Building analysis...")
    build_analysis(doc)

    print("Building conclusion...")
    build_conclusion(doc)

    print("Building references...")
    add_page_break(doc)
    build_references(doc)

    print("Building appendix A (extra figures)...")
    add_page_break(doc)
    build_appendix_a(doc)

    print("Building appendix B (interactive dashboard)...")
    add_page_break(doc)
    build_appendix_b_interactive(doc)

    print("Building appendix D (source code)...")
    add_page_break(doc)
    build_appendix_code(doc)

    doc.save(OUTPUT_FILE)
    print(f"\n[OK] Saved to: {OUTPUT_FILE}")
    print(f"     Size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")

    # GitHub instructions
    print("\n" + "=" * 60)
    print("GITHUB SETUP — do this before submitting:")
    print("=" * 60)
    print("1. Go to https://github.com/new")
    print("2. Create repo: 'sainsburys-review-analysis' (public)")
    print("3. In terminal, run:")
    print("   cd /path/to/sainsburys_project")
    print("   git init")
    print("   git add .")
    print("   git commit -m 'Initial commit: full analysis pipeline'")
    print("   git remote add origin https://github.com/YOUR_USERNAME/sainsburys-review-analysis.git")
    print("   git push -u origin main")
    print("4. Replace GITHUB_PLACEHOLDER in build_report_v2.py with your real URL")
    print("5. Re-run: python3 build_report_v2.py")
    print("=" * 60)


if __name__ == "__main__":
    main()
